from __future__ import annotations

import argparse
import json
import os
import subprocess
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = DOCS / "diagrams"
RENDERED = DIAGRAMS / "rendered"
REPORT_BUILD = DOCS / "report_build"
DOCX_OUT = DOCS / "AXIOM_Report.docx"
MD_OUT = REPORT_BUILD / "AXIOM_Report.md"
CORRECTIONS_OUT = DOCS / "AXIOM_Report_Corrections.md"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 45, 56)
MUTED = RGBColor(96, 104, 116)
FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "C9D2DE"

STUDENTS = [
    ("42210130", "Baher Mohamed"),
    ("42210128", "Shrouk Saber"),
    ("42210137", "Youssef Mohamed"),
    ("42210165", "Abanoub Attia"),
    ("42210106", "Ehab Ashraf"),
    ("42210101", "Abdelrahman Wael"),
    ("42210136", "Fady Alber"),
]

FIGURES = [
    ("Figure 3.1", "AXIOM V2 project Gantt chart", "gantt_project_plan.png"),
    ("Figure 4.1", "AXIOM V2 entity relationship diagram", "er_diagram.png"),
    ("Figure 4.2", "Relational schema mapping", "relational_schema.png"),
    ("Figure 4.3", "Code-grounded class diagram", "class_diagram.png"),
    ("Figure 4.4", "Use-case diagram", "use_case_diagram.png"),
    ("Figure 4.5", "User activity diagram", "user_activity.png"),
    ("Figure 4.6", "Admin activity diagram", "admin_activity.png"),
    ("Figure 4.7", "User sequence diagram", "user_sequence.png"),
    ("Figure 4.8", "Admin sequence diagram", "admin_sequence.png"),
    ("Figure 4.9", "Listing, booking, and subscription state diagram", "state_diagram.png"),
    ("Figure 5.1", "Software architecture and component diagram", "component_architecture.png"),
    ("Figure 6.1", "Testing and verification map", "testing_strategy.png"),
]

TABLES = [
    "Table 2.1 Competitor and platform gap analysis",
    "Table 3.1 Feasibility and estimated cost",
    "Table 3.2 Functional requirements",
    "Table 3.3 Non-functional requirements",
    "Table 3.4 User characteristics",
    "Table 4.1 Risks and mitigation strategy",
    "Table 4.2 Use-case scenarios",
    "Table 5.1 Implementation stack",
    "Table 5.2 Main implementation modules",
    "Table 5.3 User interface route map",
    "Table 6.1 Backend test coverage",
    "Table 6.2 Unit-test evidence mapped to AXIOM modules",
    "Table 6.3 Integration testing flows",
    "Table 6.4 Manual system test cases",
]


def ensure_dirs() -> None:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    RENDERED.mkdir(parents=True, exist_ok=True)
    REPORT_BUILD.mkdir(parents=True, exist_ok=True)


def d(name: str) -> Path:
    return DIAGRAMS / name


def diagram_sources() -> dict[str, str]:
    return {
        "gantt_project_plan.mmd": """---
config:
  theme: base
  themeVariables:
    primaryColor: '#E8EEF5'
    primaryTextColor: '#0B2545'
    primaryBorderColor: '#2E74B5'
    lineColor: '#7A8797'
    fontFamily: 'Arial'
    fontSize: '13px'
---
gantt
    title AXIOM V2 Graduation Project Plan
    dateFormat YYYY-MM-DD
    axisFormat %b %Y
    tickInterval 1month
    %% Source: docs/AXIOM_DEEP_OVERVIEW.md, docs/CODEBASE_STUDY.md, docs/ROADMAP.md
    section Research and Planning
    Problem research and old report review      :done, r1, 2025-10-01, 2025-11-15
    Requirements and competitor analysis        :done, r2, 2025-11-01, 2025-12-15
    section Product and UI
    UX flows and visual system                  :done, u1, 2025-12-01, 2026-01-20
    Next.js frontend pages                      :done, u2, 2026-01-10, 2026-04-30
    Dashboard and admin console                 :done, u3, 2026-03-01, 2026-05-31
    section Backend and Data
    Supabase schema and RLS                     :done, b1, 2026-01-15, 2026-03-15
    FastAPI routers and services                :done, b2, 2026-02-10, 2026-05-20
    section AI and Trust
    Ollama integration and embeddings           :done, a1, 2026-03-15, 2026-05-10
    RAG search, recommendations, fraud checks   :done, a2, 2026-04-01, 2026-05-31
    section Payments and Launch
    Stripe bookings and subscriptions           :done, p1, 2026-05-01, 2026-06-12
    Testing, documentation, deployment prep     :active, p2, 2026-05-20, 2026-06-20
""",
        "er_diagram.mmd": """---
config:
  theme: base
  layout: elk
  themeVariables:
    primaryColor: '#F4F6F9'
    primaryTextColor: '#0B2545'
    primaryBorderColor: '#2E74B5'
    lineColor: '#5E6B7A'
    fontFamily: 'Arial'
---
erDiagram
    %% Sources: docs/schema/001_v2_comprehensive_schema.sql; docs/schema/004_knowledge_chunks.sql; docs/schema/005_leads.sql; backend/sql/2026-05-15_all_new_features.sql; backend/sql/2026-05-29_payment_fees_model.sql; backend/sql/2026-05-30_owner_subscriptions.sql
    PROFILES ||--o{ LISTINGS : owns
    PROFILES ||--o{ FAVORITES : saves
    PROFILES ||--o{ LISTING_APPLICATIONS : submits
    PROFILES ||--o{ BOOKINGS : books
    PROFILES ||--o{ SUBSCRIPTIONS : has
    PROFILES ||--o{ LEADS : creates
    AGENCIES ||--o{ PROJECTS : contains
    AGENCIES ||--o{ LISTINGS : markets
    PROJECTS ||--o{ LISTINGS : includes
    NEIGHBORHOODS ||--o{ LISTINGS : locates
    LISTINGS ||--o{ HOUSEMATES : has
    LISTINGS ||--o{ LISTING_APPLICATIONS : receives
    LISTINGS ||--o{ FAVORITES : saved_as
    LISTINGS ||--o{ BOOKINGS : reserved_by
    LISTINGS ||--o{ PAYMENTS : charged_for
    LISTINGS ||--o{ KNOWLEDGE_CHUNKS : indexed_as
    LISTINGS ||--o{ LEADS : contacted_through
    LISTINGS ||--o{ VIEWINGS : visited_by
    BOOKINGS ||--o{ PAYMENTS : records
""",
        "relational_schema.puml": """@startuml
' Sources: docs/schema/001_v2_comprehensive_schema.sql; docs/schema/004_knowledge_chunks.sql; backend/sql/2026-05-15_all_new_features.sql; backend/sql/2026-05-29_payment_fees_model.sql; backend/sql/2026-05-30_owner_subscriptions.sql
skinparam backgroundColor #FFFFFF
skinparam classAttributeIconSize 0
skinparam linetype ortho
skinparam shadowing false
skinparam class {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}
hide methods

class profiles <<table>> {
  +id uuid PK auth.users(id)
  email text
  role user_role
  is_verified_seller boolean
  phone text
  lifestyle_preferences jsonb
}

class listings <<table>> {
  +id uuid PK
  owner_id uuid FK profiles.id
  agency_id uuid FK agencies.id
  project_id uuid FK projects.id
  category listing_category
  property_type property_type
  status listing_status
  price numeric
  embedding vector
}

class agencies <<table>> {
  +id uuid PK
  owner_id uuid FK profiles.id
  slug text UK
  verified boolean
}

class projects <<table>> {
  +id uuid PK
  agency_id uuid FK agencies.id
  slug text UK
  status project_status
}

class housemates <<table>> {
  +id uuid PK
  listing_id uuid FK listings.id
  name text
  lifestyle jsonb
}

class listing_applications <<table>> {
  +id uuid PK
  listing_id uuid FK listings.id
  applicant_id uuid FK profiles.id
  status application_status
  compatibility_score numeric
}

class favorites <<table>> {
  +user_id uuid PK/FK
  +listing_id uuid PK/FK
}

class bookings <<table>> {
  +id uuid PK
  listing_id uuid FK listings.id
  renter_id uuid FK profiles.id
  owner_id uuid FK profiles.id
  status text
  total_price numeric
  stripe_payment_intent_id text UK
}

class payments <<table>> {
  +id uuid PK
  user_id uuid FK profiles.id
  listing_id uuid FK listings.id
  booking_id uuid FK bookings.id
  kind text
  status text
}

class subscriptions <<table>> {
  +id uuid PK
  user_id uuid FK profiles.id UNIQUE
  plan subscription_plan
  status text
  ai_descriptions_used int
}

class knowledge_chunks <<table>> {
  +id uuid PK
  source_type text
  source_id uuid
  embedding vector
}

class leads <<table>> {
  +id uuid PK
  user_id uuid FK profiles.id
  listing_id uuid FK listings.id
  source text
}

profiles "1" --> "0..*" listings : owner_id
profiles "1" --> "0..*" agencies : owner_id
agencies "1" --> "0..*" projects : agency_id
agencies "0..1" --> "0..*" listings : agency_id
projects "0..1" --> "0..*" listings : project_id
listings "1" --> "0..*" housemates : listing_id
listings "1" --> "0..*" listing_applications : listing_id
profiles "1" --> "0..*" listing_applications : applicant_id
profiles "1" --> "0..*" favorites : user_id
listings "1" --> "0..*" favorites : listing_id
listings "1" --> "0..*" bookings : listing_id
profiles "1" --> "0..*" bookings : renter_id/owner_id
bookings "0..1" --> "0..*" payments : booking_id
profiles "1" --> "0..*" subscriptions : user_id
listings "1" --> "0..*" leads : listing_id
listings "1" --> "0..*" knowledge_chunks : source_id
@enduml
""",
        "class_diagram.puml": """@startuml
' Sources: backend/app/ai/router.py; backend/app/ai/ollama_client.py; backend/app/ai/rag.py; backend/app/ai/fraud.py; backend/app/listings/router.py; backend/app/bookings/router.py; backend/app/subscriptions/service.py; frontend/src/stores/authStore.ts; frontend/src/lib/queries.ts; frontend/src/types/api.ts
skinparam backgroundColor #FFFFFF
skinparam classAttributeIconSize 0
skinparam linetype ortho
skinparam shadowing false
skinparam class {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}

package "Frontend TypeScript" {
  class AuthStore {
    +initialize()
    +login(email,password)
    +signup(payload)
    +logout()
    +refreshProfile()
    +sendPhoneOtp(phone)
    +verifyPhoneOtp(phone, token)
  }
  class ApiClient {
    +get(path, options)
    +post(path, body)
    +put(path, body)
    +delete(path)
  }
  class QueryDefinitions {
    +listingsQueries
    +dashboardQueries
    +bookingsQueries
    +subscriptionQuery
    +ragSearchMutation
  }
  class AddListingModal {
    +validateStep()
    +generateDescription()
    +submitListing()
    +uploadImages()
  }
  class BookingModal {
    +createPaymentIntent()
    +submitPayment()
    +syncPaidBooking()
  }
}

package "FastAPI Backend" {
  class AuthRouter {
    +signup()
    +login()
    +get_me()
    +update_me()
  }
  class ListingRouter {
    +list_listings()
    +get_listing()
    +create_listing()
    +update_listing()
    +delete_listing()
    +toggle_favorite()
  }
  class BookingRouter {
    +get_fees()
    +create_payment_intent()
    +sync_payment()
    +confirm_booking()
    +refund_booking()
    +vacate_booking()
  }
  class SubscriptionService {
    +get_or_create_subscription()
    +active_listing_count()
    +increment_ai_usage()
    +ensure_listing_quota()
  }
  class StripeWebhookRouter {
    +stripe_webhook()
    -_create_booking_from_paid_intent()
    -_sync_subscription()
  }
}

package "AI Services" {
  class OllamaClient {
    +generate(prompt)
    +chat(messages)
    +generate_stream(prompt)
    +embed(text)
  }
  class RAGRetriever {
    +retrieve(query, limit)
    +build_context(chunks)
    +format_citations(chunks)
  }
  class EmbeddingsService {
    +embed_listing(listing)
    +embed_listing_chunk(listing)
    +delete_listing_chunk(listing_id)
  }
  class FraudEngine {
    +score_listing(listing)
    -price_anomaly()
    -owner_reputation()
    -llm_consistency()
  }
}

AuthStore --> ApiClient
QueryDefinitions --> ApiClient
AddListingModal --> QueryDefinitions
BookingModal --> QueryDefinitions
ListingRouter --> EmbeddingsService
ListingRouter --> FraudEngine
BookingRouter --> StripeWebhookRouter
SubscriptionService --> StripeWebhookRouter
RAGRetriever --> OllamaClient
EmbeddingsService --> OllamaClient
FraudEngine --> OllamaClient
@enduml
""",
        "use_case_diagram.puml": """@startuml
' Sources: backend/app/listings/router.py; backend/app/ai/router.py; backend/app/bookings/router.py; backend/app/subscriptions/router.py; backend/app/leads/router.py; backend/app/admin/router.py; frontend/src/app/dashboard/page.tsx; frontend/src/app/admin/dashboard/page.tsx
left to right direction
skinparam backgroundColor #FFFFFF
skinparam shadowing false
skinparam actorStyle awesome
skinparam usecase {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}

actor "Guest" as Guest
actor "User" as User
actor "Admin" as Admin
actor "AI Assistant" as AI
actor "Stripe" as Stripe
actor "Supabase" as Supabase

rectangle "AXIOM Platform" {
  usecase "Browse listings" as UC_Browse
  usecase "Search with filters" as UC_Filter
  usecase "Natural-language search" as UC_NL
  usecase "Chat with assistant" as UC_Chat
  usecase "View listing detail" as UC_Detail
  usecase "Save favorite" as UC_Favorite
  usecase "Create listing" as UC_Create
  usecase "Generate description" as UC_Desc
  usecase "Apply to shared housing" as UC_Apply
  usecase "Request viewing" as UC_Viewing
  usecase "Capture WhatsApp lead" as UC_Lead
  usecase "Book rent/shared listing" as UC_Book
  usecase "Manage dashboard" as UC_Dashboard
  usecase "Subscribe to owner plan" as UC_Subscribe
  usecase "Approve/reject listing" as UC_Approve
  usecase "Verify seller badge" as UC_Verify
  usecase "Review fraud risk" as UC_Fraud
  usecase "Manage content/entities" as UC_Content
}

Guest --> UC_Browse
Guest --> UC_Filter
Guest --> UC_NL
Guest --> UC_Chat
Guest --> UC_Detail
User --> UC_Favorite
User --> UC_Create
User --> UC_Desc
User --> UC_Apply
User --> UC_Viewing
User --> UC_Lead
User --> UC_Book
User --> UC_Dashboard
User --> UC_Subscribe
Admin --> UC_Approve
Admin --> UC_Verify
Admin --> UC_Fraud
Admin --> UC_Content
AI --> UC_NL
AI --> UC_Chat
AI --> UC_Desc
AI --> UC_Fraud
Stripe --> UC_Book
Stripe --> UC_Subscribe
Supabase --> UC_Browse
Supabase --> UC_Dashboard
@enduml
""",
        "user_activity.puml": """@startuml
' Sources: frontend/src/app/find-homes/page.tsx; frontend/src/app/property/[id]/page.tsx; frontend/src/components/booking/BookingModal.tsx; backend/app/listings/router.py; backend/app/leads/router.py; backend/app/bookings/router.py; backend/app/applications/router.py
skinparam backgroundColor #FFFFFF
skinparam activity {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}
start
:Open AXIOM web app;
if (Authenticated?) then (No)
  :Browse public listings;
  :Optionally sign up or log in through Supabase;
else (Yes)
  :Load user profile and session;
endif
:Search by filters or natural language;
:View property detail;
if (Listing category?) then (Sale)
  :Click WhatsApp contact;
  :FastAPI records lead;
  :Continue contact through wa.me link;
elseif (Shared housing)
  :Review housemates and preferences;
  if (Wants to apply?) then (Yes)
    :Submit shared-housing application;
    :Compatibility score may run;
  endif
  if (Wants to book?) then (Yes)
    :Select dates and pay booking fee through Stripe;
    :Booking is synced from webhook or fallback endpoint;
  endif
else (Rent)
  :Select move-in date and duration;
  :Pay flat booking fee through Stripe;
  :Confirm booking in dashboard;
endif
:Manage listings, favorites, bookings, applications, and profile in dashboard;
stop
@enduml
""",
        "admin_activity.puml": """@startuml
' Sources: frontend/src/app/admin/login/page.tsx; frontend/src/app/admin/dashboard/page.tsx; backend/app/admin/router.py
skinparam backgroundColor #FFFFFF
skinparam activity {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}
start
:Open admin login;
:Submit admin username and password;
if (Credentials valid?) then (Yes)
  :Receive admin JWT;
  :Open admin dashboard;
  fork
    :Review pending listings;
    if (Meets policy and trust checks?) then (Approve)
      :Set listing active;
      :Notify owner;
    else (Reject)
      :Set listing rejected with reason;
      :Notify owner;
    endif
  fork again
    :Review fraud queue;
    :Mark listing reviewed or keep under moderation;
  fork again
    :Manage users and verification badges;
  fork again
    :Manage agencies, projects, universities, and blog posts;
  end fork
  :Monitor bookings, notifications, and platform status;
else (No)
  :Show authentication error;
endif
stop
@enduml
""",
        "user_sequence.puml": """@startuml
' Sources: frontend/src/stores/authStore.ts; frontend/src/components/ai/ChatDrawer.tsx; backend/app/ai/router.py; backend/app/bookings/router.py; backend/app/stripe_webhooks/router.py
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ParticipantBackgroundColor #F4F6F9
  ParticipantBorderColor #2E74B5
  LifeLineBorderColor #7A8797
}
actor User
participant "Next.js frontend" as FE
participant "FastAPI" as API
database "Supabase Auth/Postgres" as DB
participant "Ollama/RAG" as AI
participant "Stripe" as Stripe

User -> FE: Sign in or browse
FE -> DB: Supabase Auth session
DB --> FE: JWT and user session
User -> FE: Enter natural-language search
FE -> API: POST /api/ai/search
API -> AI: Embed query and retrieve chunks
AI -> DB: hybrid_search_chunks
DB --> AI: citations and listing chunks
AI --> API: context and model response
API --> FE: parsed filters, results, citations
FE --> User: Render matched listings

User -> FE: Open rent/shared listing and book
FE -> API: POST /api/bookings/payment-intent
API -> DB: Validate listing and lock pending_payment
API -> Stripe: Create PaymentIntent
Stripe --> API: client_secret
API --> FE: booking preview and client_secret
User -> FE: Submit card payment
FE -> Stripe: confirmCardPayment()
Stripe -> API: payment_intent.succeeded webhook
API -> DB: Create booking, payment ledger, listing booked
FE -> API: POST /api/bookings/sync-payment
API --> FE: Booking detail
FE --> User: Show booking success/dashboard link
@enduml
""",
        "admin_sequence.puml": """@startuml
' Sources: frontend/src/app/admin/login/page.tsx; frontend/src/app/admin/dashboard/page.tsx; backend/app/admin/router.py; backend/app/stripe_webhooks/router.py
skinparam backgroundColor #FFFFFF
skinparam sequence {
  ParticipantBackgroundColor #F4F6F9
  ParticipantBorderColor #2E74B5
  LifeLineBorderColor #7A8797
}
actor Admin
participant "Admin dashboard" as FE
participant "Admin FastAPI router" as API
database "Supabase Postgres" as DB
participant "Notification service" as Notify
participant "Stripe webhook router" as StripeHook

Admin -> FE: Submit admin login
FE -> API: POST /api/admin/auth/login
API --> FE: Admin JWT
FE -> API: GET /api/admin/listings?status=pending
API -> DB: Query pending listings and owner profiles
DB --> API: Pending moderation queue
API --> FE: Listing rows
Admin -> FE: Approve listing
FE -> API: PUT /api/admin/listings/{id}/approve
API -> DB: Update status to active
API -> Notify: Insert owner notification
Notify -> DB: notifications row
API --> FE: Updated listing
FE --> Admin: Queue refreshes

StripeHook -> DB: Payment/subscription status updates
FE -> API: GET /api/admin/bookings
API -> DB: Query operational booking view
API --> FE: Booking/payment status data
@enduml
""",
        "state_diagram.puml": """@startuml
' Sources: backend/app/listings/router.py; backend/app/bookings/router.py; backend/app/subscriptions/plans.py; backend/app/subscriptions/lapse.py; backend/app/stripe_webhooks/router.py
skinparam backgroundColor #FFFFFF
skinparam state {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}

state "Listing lifecycle" as Listing {
  [*] --> Draft
  Draft --> Pending : submit listing
  Pending --> Active : admin approve or low fraud score
  Pending --> Rejected : admin reject
  Active --> PendingPayment : create PaymentIntent
  PendingPayment --> Active : payment canceled/refunded
  PendingPayment --> Booked : payment succeeded
  Booked --> Active : vacate or completion
  Active --> Sold : owner/admin closes sale
  Active --> Rented : owner/admin closes rental
  Active --> Paused : subscription cap exceeded
  Paused --> Active : plan restored
  Paused --> [*] : grace lapse deletion
}

state "Booking lifecycle" as Booking {
  [*] --> IntentCreated
  IntentCreated --> Paid : Stripe success
  IntentCreated --> Canceled : Stripe canceled
  Paid --> PendingConfirmation : booking row created
  PendingConfirmation --> ActiveBooking : renter confirms
  PendingConfirmation --> Refunded : refund before confirmation
  ActiveBooking --> Completed : tenant vacates or lease ends
  Refunded --> [*]
  Completed --> [*]
}

state "Subscription lifecycle" as Subscription {
  [*] --> Free
  Free --> Trialing : start trial
  Trialing --> Free : trial expires
  Trialing --> ActivePaid : checkout
  Free --> ActivePaid : checkout
  ActivePaid --> PastDue : Stripe status
  PastDue --> ActivePaid : payment recovered
  ActivePaid --> Canceled : cancel/delete
  Canceled --> Free : effective plan fallback
}
@enduml
""",
        "component_architecture.puml": """@startuml
' Sources: frontend/src/app/layout.tsx; frontend/src/providers/Providers.tsx; frontend/src/lib/api.ts; backend/app/main.py; backend/app/database.py; backend/app/ai/ollama_client.py; backend/app/stripe_client.py; railway.toml; backend/Dockerfile
skinparam backgroundColor #FFFFFF
skinparam componentStyle rectangle
skinparam shadowing false
skinparam linetype ortho
top to bottom direction
skinparam component {
  BackgroundColor #F4F6F9
  BorderColor #2E74B5
  FontColor #0B2545
}
skinparam database {
  BackgroundColor #E8EEF5
  BorderColor #1F4D78
}

actor "User/Admin Browser" as Browser
rectangle "Next.js 16 frontend" as Frontend {
  component "App Router pages" as Pages
  component "React components" as Components
  component "State and data layer" as StateLayer
  component "Direct Supabase reads" as DirectReads
}

rectangle "FastAPI backend" as Backend {
  component "Auth + profile dependencies" as Auth
  component "Listings, dashboard, applications, leads" as CoreAPI
  component "AI router + services" as AIRouter
  component "Bookings + Stripe webhooks" as Bookings
  component "Subscriptions + lapse worker" as Subs
  component "Admin router" as AdminAPI
  component "Uploads, content, universities, projects" as OtherRouters
}

rectangle "Supabase" as Supabase {
  database "Auth" as SupabaseAuth
  database "Postgres + pgvector" as Postgres
  database "Storage" as Storage
}

rectangle "External/local services" as Services {
  component "Ollama axiom-llm + embeddings" as Ollama
  component "Stripe PaymentIntents + Checkout" as Stripe
  component "Twilio Verify compatibility endpoints" as Twilio
}

Browser --> Pages
Pages --> Components
Components --> StateLayer
StateLayer --> Backend : protected /api requests with JWT
DirectReads --> Postgres : public read queries
StateLayer --> SupabaseAuth : session and OAuth/OTP
Backend --> SupabaseAuth : JWT/JWKS validation
Backend --> Postgres : service role DB operations
Backend --> Storage : signed upload URLs
AIRouter --> Ollama : generate, chat, embed
Bookings --> Stripe : PaymentIntent
Subs --> Stripe : Checkout session
Stripe --> Backend : webhook
Auth --> Twilio : compatibility OTP endpoints
@enduml
""",
        "testing_strategy.mmd": """---
config:
  theme: base
  flowchart:
    defaultRenderer: elk
  themeVariables:
    primaryColor: '#F4F6F9'
    primaryTextColor: '#0B2545'
    primaryBorderColor: '#2E74B5'
    lineColor: '#5E6B7A'
    fontFamily: 'Arial'
---
flowchart LR
    %% Sources: backend/tests/test_*.py; frontend/tsconfig.json; frontend/src/lib/queries.ts; backend/app/main.py; backend/app/stripe_webhooks/router.py
    A["Backend pytest suite<br/>123 discovered tests"] --> B["Unit and service behavior<br/>auth, listings, AI, bookings, subscriptions"]
    A --> C["API contract behavior<br/>routers, validation, protected flows"]
    D["Frontend TypeScript gate<br/>npx tsc --noEmit"] --> E["Route, store, query, and UI type safety"]
    F["Manual system tests"] --> G["Signup, login, search, AI search,<br/>listing creation, admin approval, payments, leads"]
    H["Integration scenarios"] --> I["Next.js + FastAPI + Supabase + Ollama + Stripe"]
    B --> J["Regression confidence"]
    C --> J
    E --> J
    G --> J
    I --> J
    classDef gate fill:#E8EEF5,stroke:#1F4D78,color:#0B2545,stroke-width:2px;
    classDef evidence fill:#F7FAF2,stroke:#6E8B3D,color:#1F2D16,stroke-width:1.5px;
    class A,D,F,H gate;
    class J evidence;
""",
    }


USE_CASE_SCENARIOS = [
    ["UC-A01", "Browse and search listings", "Guest/User", "The actor searches listings using filters or natural language.", "Listing search, AI search", "User enters criteria; system returns matching active listings.", "No matches; AI unavailable; invalid filters.", "Uses direct Supabase reads for public pages and FastAPI AI for natural language search."],
    ["UC-A02", "View listing detail", "Guest/User", "The actor opens a listing detail page.", "Listing search", "System loads listing data, images, owner contact, housemates when shared housing, and similar listings.", "Listing missing, inactive, or not visible.", "Shared housing is served by the same /property/[id] route."],
    ["UC-A03", "Create listing", "User", "A user publishes a rent, sale, or shared-housing listing.", "Dashboard, subscription quota", "User completes AddListingModal; backend creates a pending listing and starts fraud/embedding tasks.", "Quota exceeded, missing required fields, upload failure.", "Owner reference is owner_id."],
    ["UC-A04", "Approve or reject listing", "Admin", "Admin moderates a pending listing.", "Create listing, fraud review", "Admin reviews details and approves to active or rejects with a reason.", "Invalid admin token, listing no longer pending.", "This is the core trust gate."],
    ["UC-A05", "Generate listing description", "User, AI Assistant", "Owner asks AI to generate English/Arabic description text.", "Create listing, subscription quota", "Frontend calls /api/ai/description; backend checks quota and uses Ollama/RAG context.", "AI unavailable, quota exceeded.", "Buyer-facing AI remains free; owner description quota is plan-gated."],
    ["UC-A06", "Book rental/shared listing", "User, Stripe", "User pays a flat booking fee for a rent or shared listing.", "Listing detail, Stripe", "Backend validates listing, creates PaymentIntent, webhook creates booking and payment ledger.", "Payment canceled, listing locked, sale listing requested.", "No owner payout or Stripe Connect in current code."],
    ["UC-A07", "Apply to shared housing", "User, AI Assistant", "User applies to join a shared-housing listing.", "Shared listing detail", "Backend validates capacity and duplicate status, inserts application, and may compute compatibility.", "Not shared housing, already applied, owner applying to own listing.", "Approval increments filled spots."],
    ["UC-A08", "Capture WhatsApp lead", "User", "User requests contact through WhatsApp.", "Property detail", "Backend records or reuses a lead and returns a wa.me URL.", "No phone number or invalid source.", "Primary contact flow replacing the early inquiry model."],
    ["UC-A09", "Subscribe to owner plan", "User, Stripe", "Owner upgrades plan or starts trial.", "Pricing page", "Backend creates Stripe Checkout session or trial; webhook syncs subscription row.", "Missing Stripe config, canceled checkout, duplicate trial.", "Plan caps control listings and AI description quota."],
    ["UC-A10", "Review fraud risk", "Admin, AI Assistant", "Admin reviews listings flagged or scored by fraud logic.", "Create listing, AI fraud service", "System shows fraud queue; admin marks reviewed or moderates listing.", "AI failed open, stale listing.", "Fraud scoring supports moderation but does not replace admin control."],
]


def write_diagrams() -> None:
    ensure_dirs()
    for name, content in diagram_sources().items():
        (DIAGRAMS / name).write_text(content.strip() + "\n", encoding="utf-8")
    scenario_lines = [
        "# Use-Case Scenarios",
        "",
        "Source: backend/app/listings/router.py, backend/app/ai/router.py, backend/app/bookings/router.py, backend/app/subscriptions/router.py, backend/app/leads/router.py, backend/app/admin/router.py, frontend/src/app/dashboard/page.tsx, frontend/src/app/admin/dashboard/page.tsx",
        "",
        "| Use Case | Name | Actors | Overview | Related | User/System Action | Exceptions | Comments |",
        "| --- | --- | --- | --- | --- | --- | --- | --- |",
    ]
    for row in USE_CASE_SCENARIOS:
        scenario_lines.append("| " + " | ".join(cell.replace("|", "/") for cell in row) + " |")
    (DIAGRAMS / "use_case_scenarios.md").write_text("\n".join(scenario_lines) + "\n", encoding="utf-8")


def render_diagrams() -> None:
    ensure_dirs()
    plantuml_files = list(DIAGRAMS.glob("*.puml"))
    if plantuml_files:
        cmd = ["plantuml.cmd", "-tpng", "-o", "rendered"] + [str(p) for p in plantuml_files]
        subprocess.run(cmd, cwd=ROOT, check=True)
    mermaid_files = list(DIAGRAMS.glob("*.mmd"))
    for p in mermaid_files:
        out = RENDERED / (p.stem + ".png")
        cmd = ["cmd", "/c", "mmdc", "-i", str(p), "-o", str(out), "-b", "white", "-w", "1800", "-H", "1200", "-s", "2"]
        subprocess.run(cmd, cwd=ROOT, check=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("AXIOM Graduation Report | Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(p)


def add_para(doc, text="", style=None, bold=False, italic=False, color=None, align=None, after=None, before=None, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text)
        set_run_font(r, bold=bold, italic=italic, color=color, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        p = add_para(doc, caption, italic=True, color=MUTED, after=4, before=4, size=9.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], HEADER_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, bold=True, color=DARK_BLUE, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=9.2)
    if widths is None:
        width = int(9360 / len(headers))
        widths = [width] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)
    add_para(doc, "", after=4)
    return table


def add_figure(doc, fig_no: str, title: str, filename: str, caption: str, width=6.3):
    path = RENDERED / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
    else:
        add_para(doc, f"[Diagram image missing: {filename}]", italic=True, color=RGBColor(155, 28, 28))
    p = add_para(doc, f"{fig_no}: {title}. {caption}", italic=True, color=MUTED, size=9.5, after=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover(doc):
    for _ in range(3):
        add_para(doc, "")
    p = add_para(doc, "Faculty of Computer Science and Information Technology", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, color=DARK_BLUE, after=18)
    add_para(doc, "AXIOM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=30, color=BLUE, after=4)
    add_para(doc, "An AI-Powered Real Estate Platform for the Egyptian Market", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, color=INK, after=22)
    add_para(doc, "Submitted by:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=8)
    rows = [[str(i + 1), sid, name] for i, (sid, name) in enumerate(STUDENTS)]
    add_table(doc, ["#", "ID", "Name"], rows, widths=[800, 1900, 6660])
    add_para(doc, "A dissertation submitted in partial fulfillment of the requirements for the degree of Bachelor of Computer Science and Information Technology.", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=18)
    add_para(doc, "Supervised by:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=2)
    add_para(doc, "Dr. Bahaa Mohamed", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=DARK_BLUE, after=20)
    add_para(doc, "Egypt 2025", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()


def add_front_matter(doc):
    doc.add_heading("Committee Report", level=1)
    add_para(doc, "We certify that we have read this graduation project report as examining committee, examined the students in its content, and that in our opinion it is adequate as a project document for AXIOM.")
    add_table(doc, ["Role", "Name", "Signature", "Date"], [
        ["Chairman", "", "", ""],
        ["Supervisor", "Dr. Bahaa Mohamed", "", ""],
        ["Examiner", "", "", ""],
    ], widths=[1800, 3160, 2200, 2200])
    doc.add_page_break()

    doc.add_heading("Intellectual Property Right Declaration", level=1)
    add_para(doc, "This is to declare that the work under the supervision of Dr. Bahaa Mohamed having title AXIOM was carried out in partial fulfillment of the requirements of Bachelor of Science in Computer Science and is the sole property of Ahram Canadian University and the respective supervisor. It may be considered or used for extension, enhancement, product development, adoption for commercial or organizational usage, and related purposes only with the permission of the University and respective supervisor.")
    add_para(doc, "Names:", bold=True)
    for _, name in STUDENTS:
        add_bullet(doc, name)
    add_para(doc, "Supervisor: Dr. Bahaa Mohamed", bold=True)
    doc.add_page_break()

    doc.add_heading("Anti-Plagiarism Declaration", level=1)
    add_para(doc, "This is to declare that the above publication produced under the supervision of Dr. Bahaa Mohamed and having title AXIOM is the sole contribution of the authors. No part has been reproduced illegally or used as plagiarism. Referenced parts have been used to support the argument and have been cited properly. The authors accept responsibility for any violation if proven.")
    add_para(doc, "Names:", bold=True)
    for _, name in STUDENTS:
        add_bullet(doc, name)
    doc.add_page_break()

    doc.add_heading("Acknowledgement", level=1)
    add_para(doc, "The project team expresses sincere gratitude to Dr. Bahaa Mohamed for supervision, guidance, and continuous support throughout the project. The team also thanks the Faculty of Computer Science and Information Technology for providing the academic environment in which the project was planned, implemented, reviewed, and documented. Appreciation is extended to everyone who helped test, review, and discuss the AXIOM platform during its development.")

    doc.add_heading("Abstract", level=1)
    add_para(doc, "AXIOM is an AI-powered real estate platform designed for the Egyptian market. The platform addresses fragmented housing discovery, low trust in informal listing channels, the difficulty of finding shared housing, and the need for faster personalized search. The current implementation consists of a Next.js 16 frontend, a FastAPI backend, Supabase Auth and PostgreSQL with pgvector, local Ollama AI services, and Stripe payment integration. The system uses a unified listing model for rent, sale, and shared housing, with user ownership represented by owner_id and seller trust represented by an admin-granted verified seller badge. AXIOM provides natural-language search, chatbot support, recommendations, AI description generation, amenity validation, fraud-risk scoring, WhatsApp lead capture, viewing requests, shared-housing applications, booking payments, owner subscription plans, and an admin moderation console. This report documents the motivation, related work, planning, requirements, design, implementation, testing, limitations, and future work of AXIOM V2.")

    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "Chapter 1: Introduction",
        "Chapter 2: Background and Previous Work",
        "Chapter 3: Planning and Analysis",
        "Chapter 4: Design",
        "Chapter 5: Implementation",
        "Chapter 6: Testing",
        "Conclusion",
        "Future Work",
        "References",
        "Appendix A: Corrections from the Early Report",
    ]
    for item in toc_items:
        add_bullet(doc, item)

    doc.add_heading("List of Figures", level=1)
    for no, title, _ in FIGURES:
        add_bullet(doc, f"{no}: {title}")
    doc.add_heading("List of Tables", level=1)
    for table in TABLES:
        add_bullet(doc, table)
    doc.add_heading("List of Abbreviations and Acronyms", level=1)
    add_table(doc, ["Abbreviation", "Meaning"], [
        ["AI", "Artificial Intelligence"],
        ["API", "Application Programming Interface"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["DB", "Database"],
        ["EGP", "Egyptian Pound"],
        ["JWT", "JSON Web Token"],
        ["LLM", "Large Language Model"],
        ["NLP", "Natural Language Processing"],
        ["RAG", "Retrieval-Augmented Generation"],
        ["RLS", "Row Level Security"],
        ["SSE", "Server-Sent Events"],
        ["UI", "User Interface"],
    ], widths=[2200, 7160])
    doc.add_page_break()


def add_chapter_1(doc):
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Overview", level=2)
    add_para(doc, "AXIOM is an AI-powered real estate platform built for the Egyptian market. It supports the discovery and management of rental listings, sale listings, and shared-housing opportunities. The current implementation is a modern full-stack system that combines a Next.js frontend, a FastAPI backend, Supabase services, local Ollama AI inference, pgvector-based retrieval, and Stripe payments.")
    add_para(doc, "The platform is designed around a unified user model. A normal user can browse, save, apply, book, and publish listings. The system does not use separate buyer and seller account roles; instead, seller trust is represented by a verified seller badge granted by an administrator. This design simplifies the user experience while preserving trust controls.")

    doc.add_heading("1.2 Motivation", level=2)
    for item in [
        "Housing search in Egypt is fragmented across portals, social media groups, informal contacts, and physical agents.",
        "Users often struggle to determine whether a listing is current, accurate, safe, or fairly priced.",
        "Students, expats, and people relocating between cities need shared-housing options and compatibility signals.",
        "The 2025 rental-law context increased pressure on rental-market clarity, digital traceability, and trustworthy information.",
        "AI can reduce search effort by interpreting natural language, recommending listings, assisting owners, and supporting fraud review.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("1.3 Objective", level=2)
    add_para(doc, "The objective of AXIOM is to provide a trustworthy digital platform for housing discovery and listing management in Egypt. The project combines structured real estate data, AI-assisted discovery, shared-housing compatibility, booking support, owner subscriptions, and administrator moderation.")

    doc.add_heading("1.4 Aim", level=2)
    add_para(doc, "The aim is to reduce user search friction, improve trust in listing quality, support shared-housing decisions, and provide a maintainable software architecture that can evolve into a production real estate platform.")

    doc.add_heading("1.5 Scope", level=2)
    add_para(doc, "The functional scope includes public search, listing details, user accounts, favorites, WhatsApp lead capture, viewing requests, shared-housing applications, rent/shared-housing bookings, subscriptions, AI search, AI chat, AI recommendations, AI description generation, fraud-risk support, dashboard management, and admin moderation.")
    add_para(doc, "The technical scope includes a web frontend, REST/SSE backend APIs, Supabase authentication and PostgreSQL, vector search, local LLM inference, Stripe integration, signed uploads, and deployment preparation.")
    add_para(doc, "The user scope includes guests, authenticated users, verified sellers, administrators, and external systems such as Stripe, Supabase, and Ollama.")

    doc.add_heading("1.6 General Constraints", level=2)
    for item in [
        "The backend depends on Supabase credentials and service-role access.",
        "AI features depend on local Ollama availability and configured models.",
        "Stripe flows require configured secret keys, publishable keys, and webhook forwarding in local development.",
        "Some public frontend pages read directly from Supabase and therefore depend on correct public policies.",
        "The platform must avoid stale early-design assumptions such as separate account roles or separate shared-unit entities.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("1.7 Organization of the Dissertation", level=2)
    add_para(doc, "Chapter 2 presents background and previous work. Chapter 3 covers planning, feasibility, requirements, and analysis. Chapter 4 presents system design and diagrams. Chapter 5 explains implementation modules and user interface results. Chapter 6 discusses testing, limitations, conclusion, and future work.")


def add_chapter_2(doc):
    doc.add_heading("Chapter 2: Background and Previous Work", level=1)
    doc.add_heading("2.1 Background", level=2)
    add_para(doc, "Digital real estate platforms have become central to housing discovery, yet the Egyptian market still contains many informal search paths. Users frequently combine portal browsing, social-media posts, phone calls, and personal referrals. This creates friction and reduces confidence. AXIOM addresses this by combining structured listings, trust badges, admin approval, AI search, and localized contact behavior through WhatsApp.")
    add_para(doc, "Shared housing is a special challenge. Traditional property portals usually describe units but rarely describe resident lifestyle, compatibility, quietness, guest policies, or roommate expectations. AXIOM therefore models shared housing as a listing category with additional resident and lifestyle fields.")

    doc.add_heading("2.2 Previous Work", level=2)
    rows = [
        ["Dubizzle Egypt", "Large classifieds marketplace", "Broad inventory but limited AI matching and limited shared-housing compatibility"],
        ["Aqarmap", "Egypt-focused real estate search", "Strong local inventory but not centered on AI assistant or roommate compatibility"],
        ["Property Finder Egypt", "Regional property portal", "Professional listings but limited local shared-housing fit analysis"],
        ["Bayut Egypt", "Regional portal with Egyptian listings", "Useful structured listings but no AXIOM-style owner AI workflow"],
        ["Sakneen", "Egyptian prop-tech search", "Digital browsing but not a unified booking/subscription/AI trust workflow"],
        ["Weetas", "Regional property portal", "Curated property data but weaker student/shared-housing specialization"],
        ["Student.com", "Global student accommodation", "Student-focused but not tailored to Egyptian local listings and owners"],
        ["HousingAnywhere", "International rentals", "Useful relocation model but less localized to Egyptian trust/contact behavior"],
        ["Uniplaces", "Student accommodation platform", "Verified student housing but not an Egyptian real estate marketplace"],
        ["AmberStudent", "Global student housing", "Strong student housing focus but not local owner listing management in Egypt"],
    ]
    add_table(doc, ["Platform", "Main focus", "Gap addressed by AXIOM"], rows, widths=[1800, 2700, 4860], caption="Table 2.1 Competitor and platform gap analysis")

    doc.add_heading("2.3 Summary and Gaps", level=2)
    for item in [
        "Most competitors support listing discovery but do not combine public search, owner dashboard, admin moderation, and AI trust support in one project.",
        "Most platforms lack roommate compatibility and shared-housing resident profiles.",
        "General portals focus on filters, while AXIOM also supports natural-language search and chatbot discovery.",
        "AXIOM uses local AI services and pgvector-based retrieval, giving the project a clear technical AI layer.",
        "The current platform includes booking fees and subscription monetization, while still keeping sale listings as lead-generation flows.",
    ]:
        add_bullet(doc, item)


def add_chapter_3(doc):
    doc.add_heading("Chapter 3: Planning and Analysis", level=1)
    doc.add_heading("3.1 Planning", level=2)
    add_para(doc, "The project was planned as a full-stack web system with six major tracks: product research, interface design, database design, frontend implementation, backend implementation, AI/payments integration, and testing/deployment preparation.")

    doc.add_heading("3.1.1 Feasibility Study and Estimated Cost", level=3)
    add_table(doc, ["Area", "Feasibility", "Estimated development cost/status"], [
        ["Technical", "Feasible using Next.js, FastAPI, Supabase, Ollama, and Stripe", "Open-source stack plus cloud/service configuration"],
        ["Operational", "Feasible with admin approval and verification workflows", "Requires moderation process and policy ownership"],
        ["Economic", "Feasible through owner plans and booking fees", "Basic EGP 199/month, Pro EGP 499/month, booking fee default EGP 2000"],
        ["Schedule", "Feasible in graduation timeline with phased implementation", "Implementation completed with deployment infrastructure ready"],
        ["Risk", "Feasible with known dependency risks", "Requires configured Supabase, Stripe, and Ollama environments"],
    ], widths=[1500, 4300, 3560], caption="Table 3.1 Feasibility and estimated cost")

    doc.add_heading("3.1.2 Gantt Chart", level=3)
    add_figure(doc, "Figure 3.1", "AXIOM V2 project Gantt chart", "gantt_project_plan.png", "The plan separates research, UI design, database work, frontend and backend implementation, AI integration, payments, testing, and deployment preparation.")

    doc.add_heading("3.2 Analysis and Limitations of Existing Systems", level=2)
    add_para(doc, "Existing systems provide useful property browsing but frequently leave gaps in trust, shared-housing compatibility, owner workflow, and AI-assisted discovery. Many platforms separate user browsing from listing management and do not provide transparent moderation flows. AXIOM addresses these limitations through a unified user dashboard, administrator review, verification badges, fraud-risk scoring, and AI search.")

    doc.add_heading("3.3 Need for a New System", level=2)
    add_para(doc, "The need for AXIOM arises from the combination of market fragmentation, trust challenges, shared-housing complexity, and user expectations for fast intelligent search. A new system is justified because solving these problems requires integration across frontend UX, backend business rules, database modeling, AI services, and payment workflows.")

    doc.add_heading("3.4 Analysis of the New System", level=2)
    doc.add_heading("3.4.1 User Requirements", level=3)
    for item in [
        "Users must be able to browse listings without unnecessary friction.",
        "Users must be able to save listings, contact owners, request viewings, apply to shared housing, and book rental/shared listings.",
        "Owners must be able to create category-specific listings and manage dashboard activity.",
        "Admins must be able to approve/reject listings, verify sellers, and manage core platform entities.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3.4.2 System Requirements", level=3)
    for item in [
        "The system must validate Supabase JWTs on protected backend endpoints.",
        "The system must store unified listing records with category-specific optional fields.",
        "The system must support vector and text retrieval for AI search and chat.",
        "The system must handle Stripe PaymentIntent and Checkout webhook events.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3.4.3 Domain Requirements", level=3)
    for item in [
        "All prices are represented in EGP unless explicitly configured otherwise.",
        "Listings must support Egyptian cities, neighborhoods, compound names, and local contact behavior.",
        "Shared housing must include lifestyle and resident information.",
        "Sale listings are contact/lead-generation flows in the current code, not online payment flows.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3.4.4 Functional Requirements", level=3)
    fr_rows = [
        ["FR-01", "Auth", "Users can sign up, log in, use OAuth/OTP flows, and manage profiles."],
        ["FR-02", "Listing search", "Guests and users can browse and filter active listings."],
        ["FR-03", "AI search/chat", "Users can search naturally and chat with the AI assistant."],
        ["FR-04", "Listing creation", "Authenticated users can create rent, sale, and shared-housing listings."],
        ["FR-05", "Moderation", "Admins can approve, reject, and review risky listings."],
        ["FR-06", "Shared housing", "Users can view housemates and submit applications."],
        ["FR-07", "Bookings", "Users can pay a flat fee for rent/shared-housing bookings."],
        ["FR-08", "Subscriptions", "Owners can start trials and subscribe to Basic/Pro plans."],
        ["FR-09", "Leads", "WhatsApp lead capture records user contact actions."],
        ["FR-10", "Dashboard", "Users manage listings, favorites, bookings, applications, viewings, and profile."],
    ]
    add_table(doc, ["ID", "Requirement", "Description"], fr_rows, widths=[900, 1900, 6560], caption="Table 3.2 Functional requirements")

    doc.add_heading("3.4.5 Non-Functional Requirements", level=3)
    nfr_rows = [
        ["NFR-01", "Security", "Protected endpoints require JWT validation; admin uses separate token flow."],
        ["NFR-02", "Maintainability", "Frontend pages, API clients, backend routers, and services are separated by feature."],
        ["NFR-03", "Performance", "Public listing pages use paginated queries; AI RAG uses hybrid retrieval."],
        ["NFR-04", "Reliability", "AI and fraud paths use fail-soft/fail-open behavior where appropriate."],
        ["NFR-05", "Usability", "Category-specific forms reduce irrelevant fields and improve owner workflow."],
        ["NFR-06", "Scalability", "Supabase/PostgreSQL, pgvector, and service-layer routers support future growth."],
    ]
    add_table(doc, ["ID", "Quality", "Description"], nfr_rows, widths=[900, 1900, 6560], caption="Table 3.3 Non-functional requirements")

    doc.add_heading("3.5 Advantages of the New System", level=2)
    for item in [
        "One listing model covers rent, sale, and shared housing.",
        "AI assists discovery, owner productivity, and trust review.",
        "Admin approval improves platform quality.",
        "WhatsApp contact matches local communication behavior.",
        "Subscriptions and booking fees create a realistic monetization path.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3.6 User Characteristics", level=2)
    add_table(doc, ["User type", "Characteristics", "System support"], [
        ["Guest", "Explores available properties before committing", "Public search, detail pages, AI discovery"],
        ["Authenticated user", "Browses, saves, contacts, applies, books, and manages profile", "Dashboard, favorites, leads, viewings, applications, bookings"],
        ["Owner/verified seller", "Publishes listings and manages demand", "AddListingModal, subscription caps, AI descriptions, booking/application inboxes"],
        ["Admin", "Protects system quality and manages data", "Admin dashboard, listing moderation, fraud review, user verification, CRUD tools"],
    ], widths=[1800, 3500, 4060], caption="Table 3.4 User characteristics")


def add_chapter_4(doc):
    doc.add_heading("Chapter 4: Design", level=1)
    doc.add_heading("4.1 Design and Implementation Constraints", level=2)
    for item in [
        "The implementation must preserve the single role model of user and admin.",
        "Listing ownership must use owner_id rather than legacy ownership names.",
        "Shared housing must be represented as a listing category.",
        "Payment diagrams must reflect current code: flat booking fee, single platform Stripe account, no owner payout.",
        "AI diagrams must distinguish public AI endpoints from internal embedding and fraud services.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4.2 Assumptions and Dependencies", level=2)
    add_para(doc, "AXIOM depends on Supabase for authentication, PostgreSQL data, storage, and vector-enabled database functionality. It depends on Ollama for local LLM and embedding services. It depends on Stripe for payment and subscription events. It assumes admin moderation is part of the operating model.")

    doc.add_heading("4.3 Risks and Risk Management", level=2)
    add_table(doc, ["Risk", "Impact", "Mitigation"], [
        ["AI service unavailable", "AI search/chat/description may fail", "Use fail-soft responses and fallback search behavior"],
        ["Payment webhook not forwarded locally", "Successful payments may not immediately create bookings", "Use sync-payment fallback and configure Stripe CLI"],
        ["Schema drift", "Code may use statuses not present in a fresh database", "Verify live enum values and apply migrations before deployment"],
        ["Fake or poor-quality listings", "Reduces trust", "Use pending status, admin approval, verification, and fraud scoring"],
        ["Public direct reads misconfigured", "Pages may fail or leak data", "Review Supabase RLS policies and public query scopes"],
    ], widths=[2400, 3300, 3660], caption="Table 4.1 Risks and mitigation strategy")

    doc.add_heading("4.4 Data Design", level=2)
    doc.add_heading("4.4.1 Entity Relationship Diagram", level=3)
    add_figure(doc, "Figure 4.1", "AXIOM V2 entity relationship diagram", "er_diagram.png", "The ERD shows profiles, unified listings, shared-housing related records, bookings, payments, subscriptions, leads, knowledge chunks, and public platform entities.")
    doc.add_heading("4.4.2 Relational Schema Mapping", level=3)
    add_figure(doc, "Figure 4.2", "Relational schema mapping", "relational_schema.png", "The relational schema is derived from Supabase migrations and backend SQL files.")

    doc.add_heading("4.5 Class Diagram", level=2)
    add_figure(doc, "Figure 4.3", "Code-grounded class diagram", "class_diagram.png", "The class diagram summarizes frontend stores/components, backend routers, and AI service classes found in the code.")

    doc.add_heading("4.6 Use-Case Design", level=2)
    doc.add_heading("4.6.1 Use-Case Diagram", level=3)
    add_figure(doc, "Figure 4.4", "Use-case diagram", "use_case_diagram.png", "The diagram models public discovery, user workflows, AI services, Stripe payment/subscription use cases, and admin moderation.")
    doc.add_heading("4.6.2 Use-Case Scenarios", level=3)
    add_table(doc, ["Use Case", "Name", "Actors", "User/System Action", "Exceptions"], [[r[0], r[1], r[2], r[5], r[6]] for r in USE_CASE_SCENARIOS], widths=[900, 1700, 1500, 3560, 1700], caption="Table 4.2 Use-case scenarios")

    doc.add_heading("4.7 Activity Diagrams", level=2)
    add_figure(doc, "Figure 4.5", "User activity diagram", "user_activity.png", "The user activity flow shows browsing, authentication, AI/search discovery, property detail, payment/contact/application decisions, and dashboard management.")
    add_figure(doc, "Figure 4.6", "Admin activity diagram", "admin_activity.png", "The admin activity flow shows login, listing moderation, fraud review, verification, content management, and operational monitoring.")

    doc.add_heading("4.8 Sequence Diagrams", level=2)
    add_figure(doc, "Figure 4.7", "User sequence diagram", "user_sequence.png", "The sequence diagram follows Supabase auth, AI search/RAG, Stripe booking payment, webhook sync, and booking success.")
    add_figure(doc, "Figure 4.8", "Admin sequence diagram", "admin_sequence.png", "The admin sequence covers admin authentication, moderation queue loading, listing approval, notifications, and booking visibility.")

    doc.add_heading("4.9 State Diagram", level=2)
    add_figure(doc, "Figure 4.9", "Listing, booking, and subscription state diagram", "state_diagram.png", "The state diagram captures listing moderation and payment states, booking states, and subscription states.")


def add_chapter_5(doc):
    doc.add_heading("Chapter 5: Implementation", level=1)
    doc.add_heading("5.1 Software Architecture", level=2)
    add_figure(doc, "Figure 5.1", "Software architecture and component diagram", "component_architecture.png", "The component diagram shows the Next.js frontend, FastAPI backend, Supabase services, Ollama, Stripe, and deployment boundary.")

    doc.add_heading("5.2 Implementation Stack", level=2)
    add_table(doc, ["Layer", "Technology", "Implementation evidence"], [
        ["Frontend", "Next.js 16, TypeScript, Tailwind, Framer Motion, shadcn-style UI", "frontend/src/app, frontend/src/components"],
        ["Server state", "TanStack Query", "frontend/src/lib/queries.ts"],
        ["Client state", "Zustand", "frontend/src/stores/authStore.ts"],
        ["Backend", "FastAPI", "backend/app/main.py and feature routers"],
        ["Database", "Supabase PostgreSQL, pgvector", "docs/schema and backend/sql migrations"],
        ["Authentication", "Supabase Auth JWT", "backend/app/dependencies.py"],
        ["AI", "Ollama, RAG, embeddings, fraud scoring", "backend/app/ai"],
        ["Payments", "Stripe PaymentIntents and Checkout", "backend/app/bookings, backend/app/subscriptions"],
    ], widths=[1700, 3900, 3760], caption="Table 5.1 Implementation stack")

    doc.add_heading("5.3 Key Modules", level=2)
    add_table(doc, ["Module", "Files", "Role"], [
        ["Auth", "authStore.ts, auth/router.py, dependencies.py", "Supabase session handling, profile sync, JWT validation"],
        ["Listings", "listings/router.py, AddListingModal.tsx", "Unified listing creation, search, details, updates, favorites"],
        ["Shared housing", "applications/router.py, shared-housing components", "Applications, housemates, compatibility data"],
        ["AI/RAG", "ai/router.py, rag.py, embeddings.py, ollama_client.py", "Search, chat, recommendations, descriptions, validation, retrieval"],
        ["Fraud", "ai/fraud.py, admin/router.py", "Risk scoring and admin review support"],
        ["Bookings/Stripe", "bookings/router.py, stripe_webhooks/router.py, BookingModal.tsx", "Flat booking fee, PaymentIntent, booking sync, refund/vacate"],
        ["Subscriptions", "subscriptions/*.py, pricing/page.tsx", "Plan caps, AI quota, Stripe Checkout, lapse handling"],
        ["Leads", "leads/router.py, WhatsAppCTA.tsx", "WhatsApp lead capture and admin lead visibility"],
        ["Admin", "admin/router.py, admin/dashboard/page.tsx", "Moderation, verification, content and entity management"],
    ], widths=[1700, 3300, 4360], caption="Table 5.2 Main implementation modules")

    doc.add_heading("5.4 User Interface", level=2)
    add_table(doc, ["Route", "Purpose", "Current implementation"], [
        ["/", "Homepage", "Server page composed from public sections"],
        ["/find-homes", "Listing search", "Client page with Supabase query filters and grid/list views"],
        ["/property/[id]", "Listing detail", "Server detail page for rent, sale, and shared housing"],
        ["/shared-housing", "Shared housing search", "Client search and filtering experience"],
        ["/dashboard", "User workspace", "Tabs for listings, bookings, applications, favorites, viewings, and profile"],
        ["/pricing", "Owner plans", "Subscription state, trial, checkout, cancellation"],
        ["/admin/dashboard", "Admin console", "Live moderation, CRUD, fraud, bookings, content"],
    ], widths=[1900, 2400, 5060], caption="Table 5.3 User interface route map")

    doc.add_heading("5.5 Results and Discussion", level=2)
    add_para(doc, "The implementation demonstrates a complete modernized version of the original project idea. The platform includes all core user workflows needed for discovery, trust, shared housing, listing management, AI assistance, payments, subscriptions, and admin operations. The architecture also preserves important deployment boundaries: public browsing can read Supabase directly, while protected workflows go through FastAPI business rules.")
    add_para(doc, "The project is strongest in its integration of domain-specific requirements. Shared housing is not treated as a generic property; it has resident and lifestyle information. AI is not only a chatbot; it supports search, recommendations, owner descriptions, validation, and fraud scoring. Payments are not a future assumption; booking fees and subscriptions are implemented, while sale listing payment is intentionally excluded in the current code.")


def add_chapter_6(doc):
    doc.add_heading("Chapter 6: Testing", level=1)
    doc.add_heading("6.1 Unit Testing in AXIOM APIs and Services", level=2)
    add_para(doc, "Unit testing in AXIOM focuses on validating individual backend functions, router behaviors, service decisions, and edge cases before they are combined into complete user workflows. The project uses pytest for the FastAPI backend and verifies frontend correctness through TypeScript compilation. This testing approach is important because AXIOM depends on several moving parts: Supabase authentication, listing rules, AI fallbacks, booking payments, owner subscriptions, WhatsApp leads, and admin moderation.")
    add_para(doc, "The purpose of unit testing in the project is to confirm that each feature behaves correctly for expected inputs, handles invalid inputs safely, preserves business rules during refactoring, and avoids accidental regressions in high-risk flows such as payment creation, AI fallback behavior, seller verification, and listing approval.")
    add_figure(doc, "Figure 6.1", "Testing and verification map", "testing_strategy.png", "The map summarizes the verification evidence used for AXIOM: backend pytest tests, frontend TypeScript checking, integration scenarios, and manual system test cases.", width=6.0)

    doc.add_heading("6.1.1 Backend Test Coverage", level=3)
    add_table(doc, ["Test file", "Tests", "Coverage focus"], [
        ["test_admin.py", "6", "Admin stats, auth protection, listing update, fraud review"],
        ["test_agencies.py", "5", "Agency list/detail/projects/listings"],
        ["test_ai.py", "36", "AI search, chat, RAG, descriptions, recommendations, fraud, compatibility"],
        ["test_applications.py", "3", "Shared-housing application create/auth/category checks"],
        ["test_auth.py", "9", "Signup, login, profile, phone sync"],
        ["test_blog.py", "4", "Published blog list/detail behavior"],
        ["test_bookings.py", "5", "Flat fee behavior, sale rejection, PaymentIntent amount"],
        ["test_dashboard.py", "2", "Unified dashboard response shape"],
        ["test_leads.py", "5", "WhatsApp lead capture and duplicate handling"],
        ["test_listings.py", "20", "Listing create/list/detail/update/delete/favorite/filter/sort"],
        ["test_notifications.py", "3", "Notification list and read state"],
        ["test_ollama_client.py", "4", "Ollama client health and fallback handling"],
        ["test_projects.py", "2", "Project endpoint behavior"],
        ["test_subscriptions.py", "9", "Plan caps, trial behavior, AI quota, listing pause selection"],
        ["test_uploads.py", "2", "Signed upload URL behavior"],
        ["test_validate_amenity.py", "8", "AI amenity validation and normalization"],
        ["Total discovered", "123", "Backend pytest suite"],
    ], widths=[2300, 900, 6160], caption="Table 6.1 Backend test coverage")

    doc.add_heading("6.1.2 Unit-Test Evidence by Module", level=3)
    add_table(doc, ["Module", "Example files", "What is verified"], [
        ["Authentication", "test_auth.py", "Signup compatibility, login behavior, profile access, and phone/profile synchronization."],
        ["Listings", "test_listings.py", "Create/list/detail/update/delete behavior, filters, sorting, favorites, owner checks, and active listing visibility."],
        ["AI and RAG", "test_ai.py, test_ollama_client.py, test_validate_amenity.py", "Natural-language search, chat, recommendations, RAG fallbacks, descriptions, compatibility, fraud scoring, Ollama client handling, and amenity validation."],
        ["Shared housing", "test_applications.py", "Application creation, authentication requirements, shared-housing category validation, and duplicate/capacity rules."],
        ["Bookings and payments", "test_bookings.py", "Flat platform fee behavior, sale-listing rejection for paid bookings, and Stripe PaymentIntent amount calculation."],
        ["Subscriptions", "test_subscriptions.py", "Plan limits, trial behavior, AI description quota, and listing pause selection when caps are exceeded."],
        ["Admin and moderation", "test_admin.py", "Admin-only access, dashboard statistics, listing update behavior, and fraud review support."],
        ["Supporting modules", "test_agencies.py, test_blog.py, test_dashboard.py, test_leads.py, test_notifications.py, test_projects.py, test_uploads.py", "Public/content endpoints, dashboard aggregation, WhatsApp leads, notifications, projects, agencies, and signed uploads."],
    ], widths=[1900, 2800, 4660], caption="Table 6.2 Unit-test evidence mapped to AXIOM modules")

    doc.add_heading("6.2 Integration Testing", level=2)
    add_para(doc, "Integration testing checks that independently tested modules work together as complete product workflows. In AXIOM this means verifying that the Next.js frontend, FastAPI backend, Supabase Auth/PostgreSQL, Ollama AI services, and Stripe payment events exchange data correctly. The most important integration points are authentication, listing lifecycle, AI retrieval, booking payments, subscriptions, and dashboard aggregation.")
    add_table(doc, ["Integration flow", "Combined components", "Expected result"], [
        ["Authentication and profile load", "Next.js auth store, Supabase Auth, FastAPI JWT dependencies, profiles table", "A valid session loads the correct profile and protected routes reject unauthenticated requests."],
        ["Listing submission and moderation", "AddListingModal, FastAPI listings router, Supabase listings table, fraud/embedding tasks, admin router", "A submitted listing starts as pending, receives AI/background processing, and becomes active only after approval or permitted auto-approval."],
        ["AI search and chat", "Frontend AI UI, FastAPI AI router, RAG retriever, pgvector knowledge_chunks, Ollama", "Natural-language requests return structured results or fail softly with a usable response."],
        ["Shared-housing application", "Property detail page, applications router, listings/housemates/application tables, notifications", "A user can apply only to valid shared-housing listings with capacity and no duplicate application."],
        ["Rent/shared booking payment", "Booking modal, bookings router, Stripe PaymentIntent, stripe webhook router, bookings/payments tables", "The platform fee is calculated by the backend and payment/webhook events synchronize booking and ledger records."],
        ["Owner subscription", "Pricing page, subscriptions router, Stripe Checkout, webhook sync, subscription cap services", "Trial or paid plan state updates plan limits and AI description quota."],
        ["WhatsApp lead capture", "Property CTA, leads router, leads table, wa.me redirect", "A contact action records a lead before the user continues to WhatsApp."],
        ["Unified dashboard", "Dashboard page, dashboard router, listings/bookings/applications/favorites/viewings/profile tables", "The dashboard returns a consistent user workspace view across all relevant records."],
    ], widths=[2100, 3600, 3660], caption="Table 6.3 Integration testing flows")

    doc.add_heading("6.3 Manual System Testing", level=2)
    add_para(doc, "Manual testing is used to verify complete user-visible workflows that are difficult to prove through isolated unit tests alone. The following table is prepared in the same structure used by the reference graduation report: each case records the scenario, actions, expected result, actual result, and status. The actual-result column should be filled during the final demo rehearsal after running the frontend, backend, Supabase, Ollama, and Stripe test environment together.")
    add_table(doc, ["Test Case ID", "Scenario", "Test steps", "Expected result", "Actual result", "Status"], [
        ["TC_SIGNUP_001", "Sign up", "Open signup, enter valid name/email/password/phone data, submit.", "Account and profile are created; user can enter the app.", "To be recorded in final manual run.", "Prepared"],
        ["TC_LOGIN_002", "Login", "Open login, enter valid credentials, submit.", "Authenticated user session starts and dashboard is reachable.", "To be recorded in final manual run.", "Prepared"],
        ["TC_SEARCH_003", "Find homes", "Open /find-homes, apply city/category/price filters.", "Only matching active listings are displayed.", "To be recorded in final manual run.", "Prepared"],
        ["TC_AI_SEARCH_004", "AI search", "Enter a natural-language request such as a budget/location/preference query.", "AI endpoint returns relevant listings or a graceful fallback.", "To be recorded in final manual run.", "Prepared"],
        ["TC_CREATE_LISTING_005", "Create listing", "Open dashboard, complete the add-listing wizard for rent/sale/shared housing.", "Listing is saved with pending status and appears in owner dashboard.", "To be recorded in final manual run.", "Prepared"],
        ["TC_ADMIN_APPROVE_006", "Admin approval", "Open admin dashboard, review pending listing, approve it.", "Listing status changes to active and appears on public pages.", "To be recorded in final manual run.", "Prepared"],
        ["TC_BOOKING_007", "Booking payment", "Open a rent/shared listing, start booking, complete Stripe test payment.", "PaymentIntent and booking/payment records are synchronized.", "To be recorded in final manual run.", "Prepared"],
        ["TC_LEAD_008", "WhatsApp lead", "Click WhatsApp contact on a listing detail page.", "Lead is recorded and user continues to WhatsApp URL.", "To be recorded in final manual run.", "Prepared"],
        ["TC_SUBSCRIPTION_009", "Owner subscription", "Open pricing, start trial or checkout Basic/Pro plan.", "Subscription/trial state updates and plan limits are available.", "To be recorded in final manual run.", "Prepared"],
        ["TC_LOGOUT_010", "Logout", "Click logout from the authenticated UI.", "Session is cleared and protected pages require login again.", "To be recorded in final manual run.", "Prepared"],
    ], widths=[1050, 1300, 2550, 2150, 1650, 660], caption="Table 6.4 Manual system test cases")

    doc.add_heading("6.4 Frontend and Static Verification", level=2)
    add_para(doc, "The required frontend verification command is `npx tsc --noEmit` inside the `frontend` directory. This confirms that the TypeScript application, route components, query contracts, stores, API types, and UI props compile without type errors. It is not a replacement for browser end-to-end testing, but it is the current automated frontend gate documented for the project.")

    doc.add_heading("6.5 Limitations and Future Testing Enhancements", level=2)
    for item in [
        "Live Supabase schema must be checked for all status enum values used by the latest code.",
        "Some legacy conversation/message schema remains even though the current contact flow uses WhatsApp leads.",
        "The agency subscription stub is separate from the current owner Basic/Pro subscription system.",
        "Frontend automated browser tests are not documented as a complete suite; TypeScript is the required frontend gate and Playwright/Cypress coverage should be added for final production confidence.",
        "Manual test cases should be executed during the final demo rehearsal and updated with actual results and screenshots.",
        "Deployment infrastructure is prepared, but production deployment still requires environment and service configuration.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Conclusion", level=1)
    add_para(doc, "AXIOM V2 successfully modernizes the original real estate platform idea into a full-stack AI-assisted system. It replaces the early role-split design with a unified user model, a category-driven listing schema, and a stronger trust workflow. The implementation combines Next.js, FastAPI, Supabase, Ollama, pgvector, and Stripe to support property discovery, shared housing, AI search, owner productivity, payments, subscriptions, and admin moderation. The result is a practical and extensible foundation for a real Egyptian real estate platform.")

    doc.add_heading("Future Work", level=1)
    for item in [
        "Apply and verify all remaining live Supabase migrations, especially status enum values.",
        "Deploy the frontend and backend to production hosting with rotated keys and secure environment variables.",
        "Expand automated frontend end-to-end testing for booking, dashboard, and admin workflows.",
        "Add richer owner analytics and notification workflows.",
        "Improve multilingual Arabic/English coverage across AI prompts and UI surfaces.",
        "Extend fraud detection with more market data and reviewer feedback loops.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("References", level=1)
    refs = [
        "AXIOM V2 repository source files: frontend/src, backend/app, docs/schema, backend/sql, backend/tests.",
        "Supabase documentation, Auth, PostgreSQL, Storage, and pgvector concepts: https://supabase.com/docs",
        "FastAPI documentation: https://fastapi.tiangolo.com/",
        "Next.js documentation: https://nextjs.org/docs",
        "Stripe documentation for PaymentIntents, Checkout, and webhooks: https://docs.stripe.com/",
        "Ollama documentation: https://ollama.com/",
        "Dubizzle Egypt: https://www.dubizzle.com.eg/",
        "Aqarmap: https://aqarmap.com.eg/",
        "Property Finder Egypt: https://www.propertyfinder.eg/en/",
        "Bayut Egypt: https://www.bayut.eg/en/",
        "Sakneen: https://sakneen.com/",
        "Weetas: https://www.weetas.com/",
        "Student.com: https://www.student.com/",
        "HousingAnywhere: https://housinganywhere.com/",
        "Uniplaces: https://www.uniplaces.com/",
        "AmberStudent: https://amberstudent.com/",
        "Egypt 2025 rental-law context as discussed in the early project report and project motivation.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.add_heading("Appendix A: Corrections from the Early Report", level=1)
    add_para(doc, "The early report was used for structure and motivation only. The technical content was corrected to match AXIOM V2.")
    add_table(doc, ["Early report design", "Current report correction"], [
        ["Title and product identity", "Use AXIOM as the project title and current product name."],
        ["Separate account classes", "Use a single user/admin role model with verified seller badge."],
        ["Legacy owner naming", "Use owner_id for listing ownership."],
        ["Separate property and shared-unit entities", "Use one listings table with category values."],
        ["PHP/MySQL stack", "Use Next.js, FastAPI, Supabase PostgreSQL, pgvector, Ollama, and Stripe."],
        ["AI optional/future", "Document implemented AI search, chat, recommendations, compatibility, descriptions, validation, embeddings, and fraud scoring."],
        ["Payments future/out of scope", "Document implemented rent/shared-housing booking fee and owner subscriptions; sale listings remain contact-only."],
        ["Inquiry/contact workflow", "Document WhatsApp leads, viewing requests, applications, bookings, and dashboard tabs."],
    ], widths=[3600, 5760])


def build_docx() -> None:
    ensure_dirs()
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_chapter_1(doc)
    doc.add_page_break()
    add_chapter_2(doc)
    doc.add_page_break()
    add_chapter_3(doc)
    doc.add_page_break()
    add_chapter_4(doc)
    doc.add_page_break()
    add_chapter_5(doc)
    doc.add_page_break()
    add_chapter_6(doc)
    doc.save(DOCX_OUT)


def write_markdown_and_corrections() -> None:
    ensure_dirs()
    md = [
        "# AXIOM Graduation Report",
        "",
        "The final Word deliverable is `docs/AXIOM_Report.docx`.",
        "",
        "Generated from `docs/AXIOM_report_prompt.md`, `docs/AXIOM_DEEP_OVERVIEW.md`, and `docs/CODEBASE_STUDY.md`.",
        "",
        "## Diagram Sources",
    ]
    for source in sorted(DIAGRAMS.glob("*.*")):
        if source.name == "rendered":
            continue
        md.append(f"- `{source.relative_to(ROOT)}`")
    MD_OUT.write_text("\n".join(md) + "\n", encoding="utf-8")

    corrections = [
        "# AXIOM Report Corrections from Early Report",
        "",
        "The old early-stage report was treated as historical structure only. The following corrections were applied:",
        "",
        "1. The project title is AXIOM, not the early report title.",
        "2. The current account model is `user | admin`; the technical chapters do not model separate browsing and seller roles.",
        "3. Listing ownership uses `owner_id`.",
        "4. Shared housing is represented by `listings.category = shared_housing` with housemates and applications.",
        "5. The stack is Next.js 16, FastAPI, Supabase PostgreSQL, pgvector, Ollama, and Stripe.",
        "6. AI is implemented through seven public AI routes plus internal RAG, embedding, and fraud services.",
        "7. Payments use a flat platform-retained booking fee for rent/shared housing; sale listings are lead-generation/contact-only.",
        "8. WhatsApp leads, viewing requests, shared-housing applications, and bookings replace the early inquiry/contact flow.",
        "9. The admin system includes approval/rejection, user verification, fraud review, agencies, universities, projects, blog, and booking visibility.",
    ]
    CORRECTIONS_OUT.write_text("\n".join(corrections) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--write-diagrams", action="store_true")
    parser.add_argument("--render-diagrams", action="store_true")
    parser.add_argument("--build-docx", action="store_true")
    parser.add_argument("--write-notes", action="store_true")
    parser.add_argument("--all", action="store_true")
    args = parser.parse_args()

    if args.all or args.write_diagrams:
        write_diagrams()
    if args.all or args.render_diagrams:
        render_diagrams()
    if args.all or args.build_docx:
        build_docx()
    if args.all or args.write_notes:
        write_markdown_and_corrections()


if __name__ == "__main__":
    main()
